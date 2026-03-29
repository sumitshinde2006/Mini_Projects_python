"""Invisible Ink - Hide Encrypted Messages in Word Documents using Steganography"""

from docx import Document
from docx.shared import Pt, RGBColor
import os


def vigenere_encrypt(message, key):
    """
    Encrypt message using Vigenère cipher.
    
    Args:
        message (str): Message to encrypt
        key (str): Encryption key
    
    Returns:
        str: Encrypted message
    """
    encrypted = []
    key = key.upper()
    key_index = 0
    
    for char in message.upper():
        if char.isalpha():
            shift = ord(key[key_index % len(key)]) - ord('A')
            encrypted_char = chr((ord(char) - ord('A') + shift) % 26 + ord('A'))
            encrypted.append(encrypted_char)
            key_index += 1
        else:
            encrypted.append(char)
    
    return ''.join(encrypted)


def vigenere_decrypt(encrypted, key):
    """
    Decrypt message using Vigenère cipher.
    
    Args:
        encrypted (str): Encrypted message
        key (str): Decryption key
    
    Returns:
        str: Decrypted message
    """
    decrypted = []
    key = key.upper()
    key_index = 0
    
    for char in encrypted.upper():
        if char.isalpha():
            shift = ord(key[key_index % len(key)]) - ord('A')
            decrypted_char = chr((ord(char) - ord('A') - shift) % 26 + ord('A'))
            decrypted.append(decrypted_char)
            key_index += 1
        else:
            decrypted.append(char)
    
    return ''.join(decrypted)


def hide_message_in_docx(template_path, fake_message_path, secret_message, key, output_path):
    """
    Hide encrypted message inside Word document using steganography.
    
    Args:
        template_path (str): Path to template.docx
        fake_message_path (str): Path to fakeMessage.docx
        secret_message (str): Secret message to hide
        key (str): Encryption key for Vigenère cipher
        output_path (str): Path for output document
    """
    # Load template document
    doc = Document(template_path)
    
    # Load fake message document
    try:
        fake_doc = Document(fake_message_path)
        # Copy paragraphs from fake document
        for para in fake_doc.paragraphs:
            new_para = doc.add_paragraph(para.text)
            # Copy formatting
            new_para.style = para.style
    except FileNotFoundError:
        print(f"Warning: {fake_message_path} not found. Using template only.")
    
    # Encrypt the secret message
    encrypted_message = vigenere_encrypt(secret_message, key)
    
    # Add hidden paragraph with encrypted message
    hidden_para = doc.add_paragraph()
    hidden_run = hidden_para.add_run(encrypted_message)
    
    # Make text invisible (white color)
    hidden_run.font.color.rgb = RGBColor(255, 255, 255)  # White
    hidden_run.font.size = Pt(1)  # Very small size
    
    # Save document
    doc.save(output_path)
    print(f"✓ Hidden message saved to: {output_path}")
    print(f"  Original: {secret_message}")
    print(f"  Encrypted: {encrypted_message}")


def extract_message_from_docx(doc_path, key):
    """
    Extract and decrypt hidden message from Word document.
    
    Args:
        doc_path (str): Path to document with hidden message
        key (str): Decryption key for Vigenère cipher
    
    Returns:
        dict: Contains extracted and decrypted message
    """
    doc = Document(doc_path)
    hidden_texts = []
    
    for para in doc.paragraphs:
        for run in para.runs:
            # Check if text is white (invisible)
            if run.font.color.rgb == RGBColor(255, 255, 255):
                hidden_texts.append(run.text)
    
    result = {
        'extracted': ''.join(hidden_texts),
        'decrypted': ''
    }
    
    if result['extracted']:
        result['decrypted'] = vigenere_decrypt(result['extracted'], key)
    
    return result


def create_sample_docx(filename, content=""):
    """
    Create sample .docx file for testing.
    
    Args:
        filename (str): Output filename
        content (str): Content to add to document
    """
    doc = Document()
    
    if content:
        doc.add_paragraph(content)
    else:
        doc.add_paragraph("This is a sample document.")
    
    doc.save(filename)
    print(f"✓ Created: {filename}")


# Main execution
if __name__ == "__main__":
    print("=" * 70)
    print("Invisible Ink - Steganography with Vigenère Cipher")
    print("=" * 70)
    
    # Configuration
    KEY = "SECRET"
    MESSAGE = "ATTACK AT DAWN"
    
    # Create sample documents if they don't exist
    if not os.path.exists("template.docx"):
        print("\nCreating sample documents...")
        create_sample_docx("template.docx", "Professional Document Template")
    
    if not os.path.exists("fakeMessage.docx"):
        create_sample_docx("fakeMessage.docx", 
                          "This is a routine business document.\n\n"
                          "Please review the attached files.\n\n"
                          "Best regards")
    
    # Hide message
    print(f"\nHiding message...")
    print(f"Secret message: {MESSAGE}")
    print(f"Encryption key: {KEY}")
    
    hide_message_in_docx(
        "template.docx",
        "fakeMessage.docx",
        MESSAGE,
        KEY,
        "output_hidden.docx"
    )
    
    # Extract message
    print(f"\nExtracting message from output_hidden.docx...")
    result = extract_message_from_docx("output_hidden.docx", KEY)
    
    if result['extracted']:
        print(f"Extracted (encrypted): {result['extracted']}")
        print(f"Decrypted message: {result['decrypted']}")
    else:
        print("No hidden message found.")
    
    print("\n" + "=" * 70)
    print("Process complete!")
    print("=" * 70)
